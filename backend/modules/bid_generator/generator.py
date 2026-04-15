"""
Bid Proposal Generator — creates DOCX bid proposals from tender data.
"""
import io
from datetime import datetime, timezone
from typing import Optional
import structlog

from core.config import settings
from integrations.openai_client.client import OpenAIClient

logger = structlog.get_logger(__name__)


class BidProposalGenerator:
    def __init__(self):
        self.ai_client = OpenAIClient()

    async def generate(
        self,
        tender_data: dict,
        analysis: dict,
        profitability: dict,
        company_name: str = "Ваша компания",
        company_bin: str = "",
    ) -> bytes:
        """
        Generate a complete bid proposal as DOCX bytes.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Get AI-generated proposal text
        ai_proposal_text = await self.ai_client.generate_bid_proposal(
            tender_data=tender_data,
            analysis=analysis,
            company_name=company_name,
        )

        doc = Document()

        # Set page margins
        section = doc.sections[0]
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
        title_run.bold = True
        title_run.font.size = Pt(16)

        doc.add_paragraph()

        # Tender reference
        tender_ref = doc.add_paragraph()
        tender_ref.add_run("По тендеру: ").bold = True
        tender_ref.add_run(tender_data.get("title", ""))

        doc.add_paragraph(
            f"Дата составления: {datetime.now(timezone.utc).strftime('%d.%m.%Y')}"
        )
        doc.add_paragraph(f"Заказчик: {tender_data.get('customer_name', '')}")

        doc.add_paragraph()

        # Financial summary table
        doc.add_heading("Ценовое предложение", level=2)

        budget = tender_data.get("budget", 0)
        total_cost = profitability.get("total_cost", 0)
        profit = profitability.get("expected_profit", 0)
        margin = profitability.get("profit_margin_percent", 0)
        recommended_bid = profitability.get("recommended_bid", budget)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Показатель"
        hdr[1].text = "Сумма (₸)"

        rows = [
            ("Максимальная цена тендера", f"{budget:,.0f}"),
            ("Наша цена предложения", f"{recommended_bid:,.0f}"),
            ("Себестоимость (товар + логистика + налоги)", f"{total_cost:,.0f}"),
            ("Ожидаемая прибыль", f"{profit:,.0f}"),
            ("Маржинальность", f"{margin:.1f}%"),
        ]

        for label, value in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph()

        # AI-generated proposal content
        doc.add_heading("Техническое предложение", level=2)

        if ai_proposal_text:
            # Split by sections and add with formatting
            for line in ai_proposal_text.split("\n"):
                if line.startswith("#"):
                    level = line.count("#")
                    doc.add_heading(line.lstrip("#").strip(), level=min(level, 4))
                elif line.strip():
                    doc.add_paragraph(line.strip())
        else:
            self._add_default_proposal(doc, tender_data, analysis)

        # Delivery terms
        doc.add_paragraph()
        doc.add_heading("Условия поставки", level=2)
        lead_time = profitability.get("lead_time_days", 30)
        origin = profitability.get("origin_country", "CN")
        origin_names = {"CN": "Китай", "RU": "Россия", "KZ": "Казахстан"}

        doc.add_paragraph(f"Срок поставки: {lead_time} рабочих дней с момента заключения договора")
        doc.add_paragraph(f"Страна происхождения товара: {origin_names.get(origin, origin)}")
        doc.add_paragraph("Условия оплаты: по договорённости с заказчиком")

        # Signature block
        doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.add_run(f"\n{company_name}").bold = True
        if company_bin:
            doc.add_paragraph(f"БИН: {company_bin}")
        doc.add_paragraph("Подпись: ________________")
        doc.add_paragraph(f"Дата: {datetime.now(timezone.utc).strftime('%d.%m.%Y')}")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_default_proposal(self, doc, tender_data: dict, analysis: dict):
        """Add default proposal template when AI is not available."""
        doc.add_paragraph(
            f"Настоящим предлагаем поставку товара/услуги согласно требованиям технического задания."
        )
        doc.add_paragraph(
            f"Предмет тендера: {tender_data.get('title', '')}"
        )

        product_name = analysis.get("product_name", "")
        if product_name:
            doc.add_paragraph(f"Предлагаемый товар: {product_name}")

        analogs = analysis.get("analogs_allowed")
        if analogs:
            doc.add_paragraph(
                "Предлагаемый товар является аналогом, соответствующим техническим требованиям."
            )

        doc.add_paragraph(
            "Гарантийные обязательства: согласно действующему законодательству Республики Казахстан."
        )
