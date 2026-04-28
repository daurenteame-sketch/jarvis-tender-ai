"""
Document parser: extracts text from PDF, DOCX, HTML tender documents.

Extraction order for PDFs:
  1. pdfplumber  — fast, text-based PDFs
  2. pytesseract — OCR fallback for scanned/image-only PDFs
                   (requires: pip install pytesseract pdf2image
                    and system package: tesseract-ocr tesseract-ocr-rus)
"""
import io
from typing import Optional
import structlog

logger = structlog.get_logger(__name__)

# Minimum chars from pdfplumber before we try OCR
_PDF_MIN_CHARS = 80
# Max pages to OCR (OCR is slow; first 6 pages cover most ТЗ)
_OCR_MAX_PAGES = 6


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF bytes.

    Tries pdfplumber first (instant for text-based PDFs).
    Falls back to pytesseract OCR when the extracted text is too short,
    which indicates a scanned / image-only document.
    """
    text = _pdf_pdfplumber(content)

    if len(text.strip()) < _PDF_MIN_CHARS:
        ocr_text = _pdf_ocr(content)
        if len(ocr_text.strip()) > len(text.strip()):
            print(
                f"[document_parser] PDF: pdfplumber got {len(text)} chars — "
                f"OCR fallback got {len(ocr_text)} chars",
                flush=True,
            )
            return ocr_text
        else:
            print(
                f"[document_parser] PDF: pdfplumber got {len(text)} chars, "
                f"OCR also weak ({len(ocr_text)} chars) — using pdfplumber result",
                flush=True,
            )

    return text


def _pdf_pdfplumber(content: bytes) -> str:
    """
    Extract text from PDF using pdfplumber.
    Tries table extraction first (goszakup techspec PDFs are table-based).
    Falls back to raw text extraction.
    """
    try:
        import pdfplumber
        table_rows: list[tuple[str, str]] = []
        raw_parts: list[str] = []

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                # Try table extraction
                tables = page.extract_tables()
                for table in (tables or []):
                    for row in table:
                        if not row or len(row) < 2:
                            continue
                        # Clean cells
                        cells = [_clean_cell(c) for c in row]
                        # 2-column key-value table
                        if len(cells) == 2 and cells[0] and cells[1]:
                            table_rows.append((cells[0], cells[1]))
                        # Multi-column: join non-empty cells
                        elif any(cells):
                            joined = " | ".join(c for c in cells if c)
                            if joined.strip():
                                raw_parts.append(joined)

                # Also get plain text for non-table content
                raw = page.extract_text() or ""
                if raw.strip():
                    raw_parts.append(raw.strip())

        # If we found key-value table rows — format as clean spec
        if table_rows:
            # Drop Kazakh-language rows (keys contain Kazakh-specific characters)
            _KZ_CHARS = set("әғқңөұүіӘҒҚҢӨҰҮІ")
            ru_rows = [(k, v) for k, v in table_rows if not any(c in _KZ_CHARS for c in k)]
            # Fallback: keep all rows if no Russian rows found
            rows_to_use = ru_rows if ru_rows else table_rows
            lines = [f"{k.rstrip(': ').strip()}: {v}" for k, v in rows_to_use if k and v]
            result = "\n".join(lines)
            print(
                f"[document_parser] pdfplumber tables: {len(table_rows)} rows → {len(result)} chars",
                flush=True,
            )
            return result

        # Fallback: raw text
        result = "\n".join(raw_parts)
        print(
            f"[document_parser] pdfplumber text: {len(result)} chars",
            flush=True,
        )
        return result
    except Exception as e:
        logger.error("pdfplumber extraction failed", error=str(e))
        return ""


def _clean_cell(value) -> str:
    """Normalize a table cell value."""
    if value is None:
        return ""
    return " ".join(str(value).split()).strip()


def _pdf_ocr(content: bytes) -> str:
    """
    OCR fallback for scanned PDFs.
    Uses pdf2image to render pages, then pytesseract for OCR.
    Gracefully returns '' if either library is not installed.
    """
    try:
        import pdf2image  # type: ignore
        import pytesseract  # type: ignore
    except ImportError:
        logger.debug("OCR libraries not installed (pdf2image / pytesseract) — skipping OCR")
        return ""

    try:
        images = pdf2image.convert_from_bytes(
            content,
            dpi=200,
            first_page=1,
            last_page=_OCR_MAX_PAGES,
        )
        texts: list[str] = []
        for i, img in enumerate(images):
            # Try Russian + English; fall back to eng-only if rus lang not installed
            try:
                t = pytesseract.image_to_string(img, lang="rus+eng")
            except pytesseract.TesseractError:
                t = pytesseract.image_to_string(img, lang="eng")
            if t.strip():
                texts.append(t.strip())
                print(
                    f"[document_parser] OCR page {i+1}: {len(t)} chars",
                    flush=True,
                )
        return "\n".join(texts)
    except Exception as e:
        logger.warning("OCR failed", error=str(e))
        return ""


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error("DOCX extraction failed", error=str(e))
        return ""


def extract_text_from_bytes(content: bytes, filename: str = "") -> str:
    """Auto-detect format and extract text."""
    if not content:
        return ""

    filename_lower = filename.lower()

    # Detect by magic bytes
    if content[:4] == b"%PDF" or filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif content[:2] in (b"PK", b"\x50\x4b") or filename_lower.endswith((".docx", ".xlsx")):
        return extract_text_from_docx(content)
    else:
        # Try as plain text (Cyrillic encoding handling)
        for encoding in ("utf-8", "cp1251", "koi8-r"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")


def truncate_for_ai(text: str, max_chars: int = 8000) -> str:
    """
    Truncate text for AI analysis, keeping the most relevant parts.

    Specs are typically densest at the beginning (product name, requirements)
    and at the end (tables, final requirements).  The middle often contains
    boilerplate / legal / signature blocks that add little signal.

    Split: 75 % from the start, 25 % from the end.
    """
    if len(text) <= max_chars:
        return text
    head = int(max_chars * 0.75)   # 6 000 of 8 000
    tail = max_chars - head         # 2 000 of 8 000
    return text[:head] + "\n...[truncated]...\n" + text[-tail:]
