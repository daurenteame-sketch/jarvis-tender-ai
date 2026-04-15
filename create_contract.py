from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Поля страницы
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)

# Стиль по умолчанию
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_paragraph(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_heading(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)

# Заголовок
add_heading('ДОГОВОР КУПЛИ-ПРОДАЖИ № 01/03-2026', 14)

# Город и дата
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(12)
run1 = p.add_run('г. Алматы')
run1.font.name = 'Times New Roman'
run1.font.size = Pt(12)
# Табуляция
run2 = p.add_run('\t\t\t\t\t\t«27» марта 2026 г.')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

# Преамбула
add_paragraph(
    'Товарищество с ограниченной ответственностью «TELE SCOPE», БИН 190340006214, '
    'в лице Директора Аскендировой Сарле Канайбековны, действующей на основании Устава, '
    'именуемое в дальнейшем «Продавец», с одной стороны,',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_paragraph('и', align=WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph(
    'Товарищество с ограниченной ответственностью «Filligrand», БИН 170840001228, '
    'в лице Директора Нурмаханова Нуржана, действующего на основании Устава, '
    'именуемое в дальнейшем «Покупатель», с другой стороны,',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_paragraph(
    'совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12
)

# Раздел 1
add_heading('1. ПРЕДМЕТ ДОГОВОРА', 12)

add_paragraph(
    '1.1. Продавец обязуется передать в собственность Покупателя, а Покупатель обязуется '
    'принять и оплатить следующий товар:',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

# Таблица товара
table = doc.add_table(rows=3, cols=6)
table.style = 'Table Grid'

headers = ['№', 'Наименование товара', 'Ед. изм.', 'Кол-во', 'Цена за м²\n(с НДС)', 'Сумма\n(с НДС)']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

row1 = table.rows[1]
data = ['1', 'Панели стеновые фиброцементные\nФЦП 10х3000х1200 мм\n(278 листов)', 'м²', '1 000,8', '4 400', '4 403 520']
for i, d in enumerate(data):
    row1.cells[i].text = d
    row1.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in row1.cells[i].paragraphs[0].runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

row2 = table.rows[2]
row2.cells[1].text = 'ИТОГО:'
row2.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in row2.cells[1].paragraphs[0].runs:
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

row2.cells[5].text = '4 403 520'
row2.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in row2.cells[5].paragraphs[0].runs:
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)

doc.add_paragraph()

add_paragraph(
    'Сумма без НДС: 3 796 138 (три миллиона семьсот девяносто шесть тысяч сто тридцать восемь) тенге.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_paragraph(
    'НДС 16%: 607 382 (шестьсот семь тысяч триста восемьдесят два) тенге.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_paragraph(
    'ИТОГО С НДС: 4 403 520 (четыре миллиона четыреста три тысячи пятьсот двадцать) тенге.',
    bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY
)


# Раздел 2
add_heading('2. ЦЕНА И ПОРЯДОК ОПЛАТЫ', 12)
add_paragraph('2.1. Цена товара составляет 4 400 (четыре тысячи четыреста) тенге за 1 м², в т.ч. НДС 16%.')
add_paragraph('2.2. Общая стоимость товара составляет 4 403 520 (четыре миллиона четыреста три тысячи пятьсот двадцать) тенге, в т.ч. НДС 607 382 тенге.')
add_paragraph('2.3. Оплата производится на условиях 100% (сто процентов) предоплаты путём безналичного перечисления денежных средств на расчётный счёт Продавца до момента отгрузки товара.')
add_paragraph('2.4. Отгрузка товара производится исключительно после полного поступления денежных средств на расчётный счёт Продавца.')
add_paragraph('2.5. Датой оплаты считается дата поступления денежных средств на расчётный счёт Продавца.', space_after=12)

# Раздел 3
add_heading('3. ПОСТАВКА ТОВАРА', 12)
add_paragraph('3.1. Продавец обязуется передать товар Покупателю в течение 5 (пяти) рабочих дней с момента поступления оплаты на расчётный счёт Продавца.')
add_paragraph('3.2. Условия поставки: самовывоз со склада Продавца / доставка до адреса Покупателя (нужное подчеркнуть).')
add_paragraph('3.3. Право собственности на товар переходит к Покупателю с момента подписания товарной накладной (акта приёма-передачи).', space_after=12)

# Раздел 4
add_heading('4. ПРАВА И ОБЯЗАННОСТИ СТОРОН', 12)
add_paragraph('4.1. Продавец обязан:')
add_paragraph('     — передать товар надлежащего качества в установленные сроки;')
add_paragraph('     — передать товаросопроводительные документы (накладную, сертификат соответствия, счёт-фактуру с НДС).')
add_paragraph('4.2. Покупатель обязан:')
add_paragraph('     — принять товар и произвести его осмотр при получении;')
add_paragraph('     — оплатить товар в установленные настоящим Договором сроки;')
add_paragraph('     — при обнаружении несоответствия качества — составить акт и уведомить Продавца в течение 2 (двух) рабочих дней.', space_after=12)

# Раздел 5
add_heading('5. ОТВЕТСТВЕННОСТЬ СТОРОН', 12)
add_paragraph('5.1. За нарушение сроков оплаты Покупатель уплачивает Продавцу пеню в размере 0,1% от суммы задолженности за каждый день просрочки.')
add_paragraph('5.2. За нарушение сроков поставки Продавец уплачивает Покупателю пеню в размере 0,1% от стоимости непоставленного товара за каждый день просрочки.', space_after=12)

# Раздел 6
add_heading('6. ФОРС-МАЖОР', 12)
add_paragraph('6.1. Стороны освобождаются от ответственности за частичное или полное неисполнение обязательств, если оно явилось следствием обстоятельств непреодолимой силы.')
add_paragraph('6.2. Сторона, для которой наступили форс-мажорные обстоятельства, обязана уведомить другую сторону в течение 3 (трёх) рабочих дней.', space_after=12)

# Раздел 7
add_heading('7. РАЗРЕШЕНИЕ СПОРОВ', 12)
add_paragraph('7.1. Все споры и разногласия решаются путём переговоров.')
add_paragraph('7.2. При недостижении соглашения — в судебном порядке в соответствии с законодательством Республики Казахстан.', space_after=12)

# Раздел 8
add_heading('8. СРОК ДЕЙСТВИЯ ДОГОВОРА', 12)
add_paragraph('8.1. Договор вступает в силу с момента подписания и действует до полного исполнения Сторонами своих обязательств.', space_after=12)

# Раздел 9 — Реквизиты
add_heading('9. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН', 12)

req_table = doc.add_table(rows=12, cols=2)

left = [
    ('ПРОДАВЕЦ:', True),
    ('ТОО «TELE SCOPE»', False),
    ('БИН: 190340006214', False),
    ('г. Алматы, ул. Сапаргали Бегалина, дом 7, офис 243', False),
    ('Банк: АО «Народный Банк Казахстана»', False),
    ('ИИК: KZ96601A861009683331', False),
    ('БИК: HSBKKZKX', False),
    ('КБе: 17', False),
    ('', False),
    ('Директор ___________________________', False),
    ('Аскендирова С.К.', False),
    ('МП', False),
]

right = [
    ('ПОКУПАТЕЛЬ:', True),
    ('ТОО «Filligrand»', False),
    ('БИН: 170840001228', False),
    ('г. Шымкент, пр. Жібек жолы 75Б', False),
    ('Банк: АО «Народный Банк Казахстана»', False),
    ('ИИК: KZ916017291000004061', False),
    ('БИК: HSBKKZKX', False),
    ('КБе: 17', False),
    ('', False),
    ('Директор ___________________________', False),
    ('Нурмаханов Нуржан', False),
    ('МП', False),
]

for i in range(12):
    for j, side in enumerate([left, right]):
        cell = req_table.rows[i].cells[j]
        cell.text = side[i][0]
        for run in cell.paragraphs[0].runs:
            run.bold = side[i][1]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

doc.save('Договор_ФЦП_Filligrand.docx')
print("Файл создан: Договор_ФЦП_Filligrand.docx")
